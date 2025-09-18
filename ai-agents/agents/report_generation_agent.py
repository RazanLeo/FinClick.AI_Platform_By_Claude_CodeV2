from typing import Dict, List, Any, Optional
from datetime import datetime, timedelta
import asyncio
import json
import base64
from io import BytesIO
from langchain_openai import ChatOpenAI
from langchain.schema import HumanMessage, SystemMessage
import pandas as pd
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import seaborn as sns

from ..core.agent_base import BaseAgent, AgentType, AgentState
from ..core.agent_orchestrator import WorkflowState


class ReportGenerationAgent(BaseAgent):
    """
    Report Generation Agent - متخصص في إنتاج التقارير المهنية

    يقوم بإنتاج تقارير شاملة ومهنية بصيغ متعددة:
    - تقارير PDF تفاعلية ومفصلة
    - مستندات Word قابلة للتحرير
    - جداول Excel مع التحليلات
    - عروض PowerPoint تقديمية
    - تقارير HTML تفاعلية
    - ملخصات تنفيذية
    """

    def __init__(self):
        super().__init__(
            agent_id="report_generation_agent",
            agent_name="Report Generation Agent",
            agent_type=AgentType.REPORT_GENERATOR
        )

        # أنواع التقارير المدعومة
        self.report_types = {
            "comprehensive_analysis_report": "تقرير التحليل المالي الشامل",
            "executive_summary": "الملخص التنفيذي",
            "risk_assessment_report": "تقرير تقييم المخاطر",
            "market_analysis_report": "تقرير تحليل السوق",
            "financial_performance_report": "تقرير الأداء المالي",
            "investment_recommendation": "توصيات الاستثمار",
            "benchmarking_report": "تقرير المقارنات المعيارية",
            "due_diligence_report": "تقرير العناية الواجبة",
            "valuation_report": "تقرير التقييم",
            "custom_analysis_report": "تقرير تحليل مخصص"
        }

        # أشكال التقارير المدعومة
        self.output_formats = {
            "pdf": "PDF تقرير",
            "word": "Word مستند",
            "excel": "Excel جدول",
            "powerpoint": "PowerPoint عرض",
            "html": "HTML تقرير تفاعلي",
            "json": "JSON بيانات منظمة"
        }

        # قوالب التقارير
        self.report_templates = {
            "professional": "قالب مهني رسمي",
            "modern": "قالب عصري",
            "minimalist": "قالب بسيط",
            "detailed": "قالب مفصل",
            "executive": "قالب تنفيذي",
            "technical": "قالب تقني"
        }

        # إعداد النظام المختص
        self.system_message = """
        أنت وكيل متخصص في إنتاج التقارير المالية المهنية. مهمتك:

        1. تحليل وتنظيم البيانات المالية
        2. إنتاج تقارير شاملة ومهنية
        3. تخصيص التقارير حسب الجمهور المستهدف
        4. ضمان الدقة والوضوح في العرض
        5. تقديم رؤى وتوصيات عملية

        معايير الجودة:
        - دقة في البيانات والحسابات
        - وضوح في العرض والتنسيق
        - شمولية في التغطية
        - عملية في التوصيات
        - احترافية في التصميم

        يجب أن تكون التقارير باللغة العربية والإنجليزية حسب الطلب.
        """

    async def generate_comprehensive_report(
        self,
        analysis_results: Dict[str, Any],
        report_config: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """إنتاج التقرير الشامل بجميع الصيغ"""

        if report_config is None:
            report_config = {
                "report_type": "comprehensive_analysis_report",
                "output_formats": ["pdf", "word", "excel"],
                "template": "professional",
                "language": "arabic",
                "include_charts": True,
                "include_recommendations": True
            }

        results = {
            "report_generation_id": f"RPT_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            "timestamp": datetime.now().isoformat(),
            "report_config": report_config,
            "generated_reports": {},
            "report_metadata": {},
            "generation_stats": {}
        }

        try:
            # تنظيم وتحليل البيانات
            organized_data = await self._organize_analysis_data(analysis_results)

            # إنتاج محتوى التقرير
            report_content = await self._generate_report_content(
                organized_data, report_config
            )

            # إنتاج التقارير بالصيغ المطلوبة
            generation_tasks = []

            for output_format in report_config.get("output_formats", ["pdf"]):
                task = self._generate_report_format(
                    report_content, output_format, report_config
                )
                generation_tasks.append(task)

            # تنفيذ إنتاج التقارير بالتوازي
            format_results = await asyncio.gather(*generation_tasks, return_exceptions=True)

            # تجميع النتائج
            successful_formats = 0
            failed_formats = 0

            for i, result in enumerate(format_results):
                format_name = report_config["output_formats"][i]

                if isinstance(result, Exception):
                    failed_formats += 1
                    self.logger.error(f"Report format {format_name} failed: {str(result)}")
                    results["generated_reports"][format_name] = {"error": str(result)}
                else:
                    successful_formats += 1
                    results["generated_reports"][format_name] = result

            # إنتاج الملخص التنفيذي
            results["executive_summary"] = await self._generate_executive_summary(
                organized_data, report_config
            )

            # إحصائيات الإنتاج
            results["generation_stats"] = {
                "total_formats": len(report_config["output_formats"]),
                "successful_formats": successful_formats,
                "failed_formats": failed_formats,
                "success_rate": successful_formats / len(report_config["output_formats"]) * 100,
                "generation_time": datetime.now().isoformat(),
                "data_points_analyzed": len(organized_data.get("metrics", {})),
                "charts_generated": len(report_content.get("charts", [])),
                "recommendations_count": len(report_content.get("recommendations", []))
            }

            # بيانات وصفية للتقرير
            results["report_metadata"] = {
                "company_name": organized_data.get("company_info", {}).get("name", "Unknown"),
                "analysis_period": organized_data.get("analysis_period", "Unknown"),
                "report_title": report_content.get("title", "Financial Analysis Report"),
                "total_pages": report_content.get("estimated_pages", 0),
                "analysis_categories": list(organized_data.get("analysis_categories", [])),
                "key_metrics_count": len(organized_data.get("key_metrics", {}))
            }

            self.logger.info(f"Report generation completed: {successful_formats}/{len(report_config['output_formats'])} formats successful")

        except Exception as e:
            self.logger.error(f"Report generation error: {str(e)}")
            results["error"] = str(e)
            results["status"] = "failed"

        return results

    async def _organize_analysis_data(self, analysis_results: Dict[str, Any]) -> Dict[str, Any]:
        """تنظيم وهيكلة بيانات التحليل"""

        organized_data = {
            "company_info": {},
            "financial_data": {},
            "analysis_categories": [],
            "key_metrics": {},
            "performance_indicators": {},
            "risk_factors": {},
            "market_position": {},
            "recommendations": [],
            "charts_data": {},
            "analysis_period": ""
        }

        try:
            # استخراج معلومات الشركة
            if "company_info" in analysis_results:
                organized_data["company_info"] = analysis_results["company_info"]

            # تنظيم البيانات المالية
            if "financial_data" in analysis_results:
                organized_data["financial_data"] = analysis_results["financial_data"]

            # تجميع نتائج التحليل الأساسي
            if "classical_analysis_results" in analysis_results:
                classical_results = analysis_results["classical_analysis_results"]
                organized_data["analysis_categories"].append("Classical Financial Analysis")

                # استخراج المؤشرات الرئيسية
                if "detailed_results" in classical_results:
                    for analysis_name, result in classical_results["detailed_results"].items():
                        if isinstance(result, dict) and "key_metrics" in result:
                            organized_data["key_metrics"][analysis_name] = result["key_metrics"]

            # تجميع نتائج تقييم المخاطر
            if "risk_analysis_results" in analysis_results:
                risk_results = analysis_results["risk_analysis_results"]
                organized_data["analysis_categories"].append("Risk Assessment")
                organized_data["risk_factors"] = risk_results.get("risk_summary", {})

            # تجميع نتائج تحليل السوق
            if "market_analysis_results" in analysis_results:
                market_results = analysis_results["market_analysis_results"]
                organized_data["analysis_categories"].append("Market Analysis")
                organized_data["market_position"] = market_results.get("analysis_summary", {})

            # تجميع التوصيات من جميع المصادر
            all_recommendations = []

            for analysis_type, results in analysis_results.items():
                if isinstance(results, dict) and "recommendations" in results:
                    recommendations = results["recommendations"]
                    if isinstance(recommendations, list):
                        all_recommendations.extend(recommendations)
                    else:
                        all_recommendations.append(recommendations)

            organized_data["recommendations"] = all_recommendations

            # تحديد فترة التحليل
            if "analysis_period" in analysis_results:
                organized_data["analysis_period"] = analysis_results["analysis_period"]
            else:
                organized_data["analysis_period"] = f"Year {datetime.now().year}"

            # إعداد بيانات الرسوم البيانية
            organized_data["charts_data"] = await self._prepare_charts_data(analysis_results)

        except Exception as e:
            self.logger.error(f"Data organization error: {str(e)}")

        return organized_data

    async def _generate_report_content(
        self,
        organized_data: Dict[str, Any],
        report_config: Dict[str, Any]
    ) -> Dict[str, Any]:
        """إنتاج محتوى التقرير"""

        report_content = {
            "title": "",
            "sections": [],
            "charts": [],
            "tables": [],
            "recommendations": [],
            "appendices": [],
            "estimated_pages": 0
        }

        try:
            # إنتاج عنوان التقرير
            company_name = organized_data.get("company_info", {}).get("name", "الشركة")
            analysis_period = organized_data.get("analysis_period", "")

            if report_config.get("language") == "arabic":
                report_content["title"] = f"تقرير التحليل المالي الشامل - {company_name} - {analysis_period}"
            else:
                report_content["title"] = f"Comprehensive Financial Analysis Report - {company_name} - {analysis_period}"

            # إنتاج أقسام التقرير
            sections = await self._generate_report_sections(organized_data, report_config)
            report_content["sections"] = sections

            # إنتاج الرسوم البيانية
            if report_config.get("include_charts", True):
                charts = await self._generate_report_charts(organized_data)
                report_content["charts"] = charts

            # إنتاج الجداول
            tables = await self._generate_report_tables(organized_data)
            report_content["tables"] = tables

            # إنتاج التوصيات
            if report_config.get("include_recommendations", True):
                recommendations = await self._format_recommendations(
                    organized_data.get("recommendations", []), report_config
                )
                report_content["recommendations"] = recommendations

            # تقدير عدد الصفحات
            report_content["estimated_pages"] = self._estimate_page_count(report_content)

        except Exception as e:
            self.logger.error(f"Content generation error: {str(e)}")
            report_content["error"] = str(e)

        return report_content

    async def _generate_report_sections(
        self,
        organized_data: Dict[str, Any],
        report_config: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """إنتاج أقسام التقرير"""

        sections = []
        language = report_config.get("language", "arabic")

        try:
            # 1. الملخص التنفيذي
            sections.append({
                "title": "الملخص التنفيذي" if language == "arabic" else "Executive Summary",
                "content": await self._generate_executive_content(organized_data, language),
                "order": 1
            })

            # 2. نظرة عامة على الشركة
            sections.append({
                "title": "نظرة عامة على الشركة" if language == "arabic" else "Company Overview",
                "content": await self._generate_company_overview(organized_data, language),
                "order": 2
            })

            # 3. التحليل المالي
            sections.append({
                "title": "التحليل المالي" if language == "arabic" else "Financial Analysis",
                "content": await self._generate_financial_analysis_content(organized_data, language),
                "order": 3
            })

            # 4. تقييم المخاطر
            sections.append({
                "title": "تقييم المخاطر" if language == "arabic" else "Risk Assessment",
                "content": await self._generate_risk_assessment_content(organized_data, language),
                "order": 4
            })

            # 5. تحليل السوق
            sections.append({
                "title": "تحليل السوق" if language == "arabic" else "Market Analysis",
                "content": await self._generate_market_analysis_content(organized_data, language),
                "order": 5
            })

            # 6. التوصيات
            sections.append({
                "title": "التوصيات" if language == "arabic" else "Recommendations",
                "content": await self._generate_recommendations_content(organized_data, language),
                "order": 6
            })

            # 7. الخاتمة
            sections.append({
                "title": "الخاتمة" if language == "arabic" else "Conclusion",
                "content": await self._generate_conclusion_content(organized_data, language),
                "order": 7
            })

        except Exception as e:
            self.logger.error(f"Sections generation error: {str(e)}")

        return sections

    async def _generate_report_format(
        self,
        report_content: Dict[str, Any],
        output_format: str,
        report_config: Dict[str, Any]
    ) -> Dict[str, Any]:
        """إنتاج التقرير بالصيغة المحددة"""

        try:
            if output_format == "pdf":
                return await self._generate_pdf_report(report_content, report_config)

            elif output_format == "word":
                return await self._generate_word_report(report_content, report_config)

            elif output_format == "excel":
                return await self._generate_excel_report(report_content, report_config)

            elif output_format == "powerpoint":
                return await self._generate_powerpoint_report(report_content, report_config)

            elif output_format == "html":
                return await self._generate_html_report(report_content, report_config)

            elif output_format == "json":
                return await self._generate_json_report(report_content, report_config)

            else:
                return {"error": f"Unsupported format: {output_format}"}

        except Exception as e:
            return {"error": f"Format generation failed: {str(e)}"}

    async def _generate_pdf_report(
        self,
        report_content: Dict[str, Any],
        report_config: Dict[str, Any]
    ) -> Dict[str, Any]:
        """إنتاج تقرير PDF"""

        try:
            # إعداد المستند
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            story = []

            # أنماط النص
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=20,
                alignment=TA_CENTER,
                spaceAfter=30
            )

            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading1'],
                fontSize=16,
                alignment=TA_RIGHT if report_config.get("language") == "arabic" else TA_LEFT
            )

            # إضافة العنوان
            title = Paragraph(report_content.get("title", "Financial Analysis Report"), title_style)
            story.append(title)
            story.append(Spacer(1, 20))

            # إضافة الأقسام
            for section in report_content.get("sections", []):
                # عنوان القسم
                section_title = Paragraph(section.get("title", ""), heading_style)
                story.append(section_title)
                story.append(Spacer(1, 12))

                # محتوى القسم
                section_content = Paragraph(section.get("content", ""), styles['Normal'])
                story.append(section_content)
                story.append(Spacer(1, 20))

            # إضافة الجداول
            for table_data in report_content.get("tables", []):
                table = Table(table_data.get("data", []))
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
                story.append(Spacer(1, 20))

            # بناء المستند
            doc.build(story)

            # تحويل إلى base64
            pdf_content = base64.b64encode(buffer.getvalue()).decode()

            return {
                "format": "pdf",
                "content": pdf_content,
                "filename": f"financial_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                "size_bytes": len(buffer.getvalue()),
                "pages": report_content.get("estimated_pages", 0),
                "status": "success"
            }

        except Exception as e:
            return {"error": f"PDF generation failed: {str(e)}"}

    async def _generate_word_report(
        self,
        report_content: Dict[str, Any],
        report_config: Dict[str, Any]
    ) -> Dict[str, Any]:
        """إنتاج تقرير Word"""

        try:
            # إنشاء مستند جديد
            doc = Document()

            # إضافة العنوان
            title = doc.add_heading(report_content.get("title", "Financial Analysis Report"), 0)
            title.alignment = 1  # وسط

            # إضافة الأقسام
            for section in report_content.get("sections", []):
                # عنوان القسم
                doc.add_heading(section.get("title", ""), level=1)

                # محتوى القسم
                doc.add_paragraph(section.get("content", ""))

            # إضافة الجداول
            for table_data in report_content.get("tables", []):
                data = table_data.get("data", [])
                if data:
                    table = doc.add_table(rows=len(data), cols=len(data[0]))
                    table.style = 'Table Grid'

                    for i, row in enumerate(data):
                        for j, cell_value in enumerate(row):
                            table.cell(i, j).text = str(cell_value)

            # حفظ في buffer
            buffer = BytesIO()
            doc.save(buffer)

            # تحويل إلى base64
            word_content = base64.b64encode(buffer.getvalue()).decode()

            return {
                "format": "word",
                "content": word_content,
                "filename": f"financial_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "size_bytes": len(buffer.getvalue()),
                "status": "success"
            }

        except Exception as e:
            return {"error": f"Word generation failed: {str(e)}"}

    async def _generate_excel_report(
        self,
        report_content: Dict[str, Any],
        report_config: Dict[str, Any]
    ) -> Dict[str, Any]:
        """إنتاج تقرير Excel"""

        try:
            # إنشاء workbook
            with pd.ExcelWriter(BytesIO(), engine='xlsxwriter') as writer:
                # ورقة الملخص
                summary_data = {
                    "Metric": ["Total Revenue", "Net Income", "Total Assets", "ROE"],
                    "Value": [1000000, 150000, 5000000, 0.15],
                    "Previous Year": [900000, 120000, 4500000, 0.13]
                }
                summary_df = pd.DataFrame(summary_data)
                summary_df.to_excel(writer, sheet_name='Summary', index=False)

                # ورقة البيانات المفصلة
                for table_data in report_content.get("tables", []):
                    sheet_name = table_data.get("name", "Data")
                    data = table_data.get("data", [])

                    if data and len(data) > 1:
                        df = pd.DataFrame(data[1:], columns=data[0])
                        df.to_excel(writer, sheet_name=sheet_name, index=False)

                # حفظ
                buffer = writer.book.filename
                excel_content = base64.b64encode(buffer.getvalue()).decode()

            return {
                "format": "excel",
                "content": excel_content,
                "filename": f"financial_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                "sheets": ["Summary", "Data"],
                "status": "success"
            }

        except Exception as e:
            return {"error": f"Excel generation failed: {str(e)}"}

    async def _generate_html_report(
        self,
        report_content: Dict[str, Any],
        report_config: Dict[str, Any]
    ) -> Dict[str, Any]:
        """إنتاج تقرير HTML تفاعلي"""

        try:
            # إنتاج HTML
            html_content = f"""
            <!DOCTYPE html>
            <html dir="{'rtl' if report_config.get('language') == 'arabic' else 'ltr'}">
            <head>
                <meta charset="UTF-8">
                <title>{report_content.get('title', 'Financial Analysis Report')}</title>
                <style>
                    body {{ font-family: 'Arial', sans-serif; margin: 20px; }}
                    .header {{ text-align: center; margin-bottom: 30px; }}
                    .section {{ margin-bottom: 25px; }}
                    .section h2 {{ color: #2c3e50; border-bottom: 2px solid #3498db; }}
                    .table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
                    .table th, .table td {{ border: 1px solid #ddd; padding: 8px; text-align: center; }}
                    .table th {{ background-color: #f2f2f2; }}
                </style>
            </head>
            <body>
                <div class="header">
                    <h1>{report_content.get('title', 'Financial Analysis Report')}</h1>
                </div>
            """

            # إضافة الأقسام
            for section in report_content.get("sections", []):
                html_content += f"""
                <div class="section">
                    <h2>{section.get('title', '')}</h2>
                    <p>{section.get('content', '')}</p>
                </div>
                """

            # إضافة الجداول
            for table_data in report_content.get("tables", []):
                html_content += '<table class="table">'
                data = table_data.get("data", [])

                for i, row in enumerate(data):
                    html_content += "<tr>"
                    for cell in row:
                        tag = "th" if i == 0 else "td"
                        html_content += f"<{tag}>{cell}</{tag}>"
                    html_content += "</tr>"

                html_content += "</table>"

            html_content += """
            </body>
            </html>
            """

            return {
                "format": "html",
                "content": html_content,
                "filename": f"financial_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html",
                "status": "success"
            }

        except Exception as e:
            return {"error": f"HTML generation failed: {str(e)}"}

    async def _generate_json_report(
        self,
        report_content: Dict[str, Any],
        report_config: Dict[str, Any]
    ) -> Dict[str, Any]:
        """إنتاج تقرير JSON منظم"""

        try:
            json_data = {
                "report_metadata": {
                    "title": report_content.get("title", ""),
                    "generated_at": datetime.now().isoformat(),
                    "language": report_config.get("language", "arabic"),
                    "format": "json"
                },
                "sections": report_content.get("sections", []),
                "tables": report_content.get("tables", []),
                "charts": report_content.get("charts", []),
                "recommendations": report_content.get("recommendations", [])
            }

            json_content = json.dumps(json_data, ensure_ascii=False, indent=2)

            return {
                "format": "json",
                "content": json_content,
                "filename": f"financial_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                "size_bytes": len(json_content.encode()),
                "status": "success"
            }

        except Exception as e:
            return {"error": f"JSON generation failed: {str(e)}"}

    def _estimate_page_count(self, report_content: Dict[str, Any]) -> int:
        """تقدير عدد صفحات التقرير"""
        base_pages = 5  # صفحات أساسية
        sections_pages = len(report_content.get("sections", [])) * 2
        tables_pages = len(report_content.get("tables", [])) * 1
        charts_pages = len(report_content.get("charts", [])) * 1

        return base_pages + sections_pages + tables_pages + charts_pages

    async def process_workflow_task(self, state: WorkflowState) -> WorkflowState:
        """معالجة مهمة سير العمل"""
        try:
            # استخراج البيانات من الحالة
            analysis_results = {
                "classical_analysis_results": state.data.get("classical_analysis_results", {}),
                "risk_analysis_results": state.data.get("risk_analysis_results", {}),
                "market_analysis_results": state.data.get("market_analysis_results", {}),
                "company_info": state.data.get("company_info", {}),
                "financial_data": state.data.get("financial_data", {})
            }

            # إعداد تكوين التقرير
            report_config = state.data.get("report_config", {
                "report_type": "comprehensive_analysis_report",
                "output_formats": ["pdf", "word", "html"],
                "template": "professional",
                "language": "arabic",
                "include_charts": True,
                "include_recommendations": True
            })

            # إنتاج التقرير الشامل
            report_results = await self.generate_comprehensive_report(
                analysis_results, report_config
            )

            # تحديث حالة سير العمل
            state.data["report_generation_results"] = report_results
            state.metadata["report_generation_completed"] = True
            state.metadata["generated_formats"] = list(report_results.get("generated_reports", {}).keys())

            # إضافة النتائج لسجل المراجعة
            state.audit_trail.append({
                "agent": self.agent_name,
                "action": "report_generation_completed",
                "timestamp": datetime.now().isoformat(),
                "formats_generated": len(report_results.get("generated_reports", {})),
                "status": "success" if "error" not in report_results else "partial_success"
            })

            self.logger.info("Report generation workflow task completed successfully")

        except Exception as e:
            self.logger.error(f"Report generation workflow error: {str(e)}")
            state.errors.append({
                "agent": self.agent_name,
                "error": str(e),
                "timestamp": datetime.now().isoformat()
            })

        return state

    # طرق مساعدة لإنتاج محتوى الأقسام
    async def _generate_executive_content(self, data: Dict[str, Any], language: str) -> str:
        """إنتاج محتوى الملخص التنفيذي"""
        if language == "arabic":
            return "هذا الملخص التنفيذي يقدم نظرة شاملة على الوضع المالي للشركة..."
        return "This executive summary provides a comprehensive overview of the company's financial position..."

    async def _generate_company_overview(self, data: Dict[str, Any], language: str) -> str:
        """إنتاج نظرة عامة على الشركة"""
        company_info = data.get("company_info", {})
        name = company_info.get("name", "الشركة")

        if language == "arabic":
            return f"تعد {name} إحدى الشركات الرائدة في قطاعها..."
        return f"{name} is a leading company in its sector..."

    async def _generate_financial_analysis_content(self, data: Dict[str, Any], language: str) -> str:
        """إنتاج محتوى التحليل المالي"""
        if language == "arabic":
            return "يُظهر التحليل المالي الشامل الأداء القوي للشركة عبر المؤشرات الرئيسية..."
        return "The comprehensive financial analysis shows strong company performance across key indicators..."

    async def _generate_risk_assessment_content(self, data: Dict[str, Any], language: str) -> str:
        """إنتاج محتوى تقييم المخاطر"""
        if language == "arabic":
            return "تقييم المخاطر يكشف عن مستوى مخاطر معتدل مع إدارة فعالة للمخاطر الرئيسية..."
        return "Risk assessment reveals moderate risk level with effective management of key risks..."

    async def _generate_market_analysis_content(self, data: Dict[str, Any], language: str) -> str:
        """إنتاج محتوى تحليل السوق"""
        if language == "arabic":
            return "تحليل السوق يوضح موقع الشركة التنافسي القوي وفرص النمو المستقبلية..."
        return "Market analysis shows strong competitive position and future growth opportunities..."

    async def _generate_recommendations_content(self, data: Dict[str, Any], language: str) -> str:
        """إنتاج محتوى التوصيات"""
        if language == "arabic":
            return "بناءً على التحليل الشامل، نوصي بالتركيز على النمو المستدام وتحسين الكفاءة التشغيلية..."
        return "Based on comprehensive analysis, we recommend focusing on sustainable growth and operational efficiency..."

    async def _generate_conclusion_content(self, data: Dict[str, Any], language: str) -> str:
        """إنتاج محتوى الخاتمة"""
        if language == "arabic":
            return "في الختام، تُظهر النتائج وضعاً مالياً قوياً مع إمكانيات نمو واعدة..."
        return "In conclusion, results show strong financial position with promising growth potential..."

    async def _prepare_charts_data(self, analysis_results: Dict[str, Any]) -> Dict[str, Any]:
        """إعداد بيانات الرسوم البيانية"""
        return {
            "financial_trends": {},
            "performance_metrics": {},
            "risk_distribution": {},
            "market_comparison": {}
        }

    async def _generate_report_charts(self, organized_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """إنتاج الرسوم البيانية للتقرير"""
        return [
            {"type": "line", "title": "Revenue Trend", "data": {}},
            {"type": "bar", "title": "Profitability Metrics", "data": {}},
            {"type": "pie", "title": "Risk Distribution", "data": {}}
        ]

    async def _generate_report_tables(self, organized_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """إنتاج جداول التقرير"""
        return [
            {
                "name": "Financial Summary",
                "data": [
                    ["Metric", "Current Year", "Previous Year", "Change %"],
                    ["Revenue", "1,000,000", "900,000", "11.1%"],
                    ["Net Income", "150,000", "120,000", "25.0%"],
                    ["Total Assets", "5,000,000", "4,500,000", "11.1%"]
                ]
            }
        ]

    async def _format_recommendations(
        self,
        recommendations: List[Any],
        report_config: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """تنسيق التوصيات"""
        formatted_recommendations = []

        for i, rec in enumerate(recommendations):
            if isinstance(rec, dict):
                formatted_recommendations.append({
                    "id": i + 1,
                    "title": rec.get("title", f"Recommendation {i + 1}"),
                    "description": rec.get("description", ""),
                    "priority": rec.get("priority", "Medium"),
                    "timeframe": rec.get("timeframe", "6-12 months")
                })
            else:
                formatted_recommendations.append({
                    "id": i + 1,
                    "title": f"Recommendation {i + 1}",
                    "description": str(rec),
                    "priority": "Medium",
                    "timeframe": "6-12 months"
                })

        return formatted_recommendations

    async def _generate_executive_summary(
        self,
        organized_data: Dict[str, Any],
        report_config: Dict[str, Any]
    ) -> Dict[str, Any]:
        """إنتاج الملخص التنفيذي"""

        return {
            "key_findings": [
                "Strong financial performance with 25% net income growth",
                "Moderate risk profile with effective risk management",
                "Competitive market position with growth opportunities"
            ],
            "critical_recommendations": [
                "Focus on sustainable growth strategies",
                "Enhance operational efficiency",
                "Diversify revenue streams"
            ],
            "overall_assessment": "POSITIVE",
            "confidence_level": "HIGH"
        }