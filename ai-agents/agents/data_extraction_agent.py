"""
Data Extraction Agent
وكيل استخراج البيانات

This agent specializes in extracting financial data from various document formats
including PDFs, Excel files, Word documents, and scanned images using OCR.
"""

from typing import Dict, Any, List, Optional
import asyncio
import json
import re
from datetime import datetime
import base64
import io

# Document processing imports
try:
    import PyPDF2
    import openpyxl
    import pandas as pd
    from PIL import Image
    import pytesseract
    from docx import Document
except ImportError:
    print("Warning: Some document processing libraries not installed")

from ..core.agent_base import FinancialAgent, AgentType, AgentTask
from langchain_core.prompts import ChatPromptTemplate


class DataExtractionAgent(FinancialAgent):
    """
    Specialized agent for extracting financial data from documents
    وكيل متخصص في استخراج البيانات المالية من المستندات
    """

    def __init__(self, agent_id: str, agent_name_ar: str, agent_name_en: str):
        super().__init__(
            agent_id=agent_id,
            agent_name=f"{agent_name_ar} | {agent_name_en}",
            agent_type=AgentType.DATA_EXTRACTION
        )

        # Extraction patterns and templates
        self.financial_patterns = self._initialize_extraction_patterns()
        self.document_processors = self._initialize_document_processors()

    def _initialize_capabilities(self) -> None:
        """Initialize data extraction capabilities"""
        super()._initialize_capabilities()

        extraction_capabilities = [
            "pdf_extraction",
            "excel_extraction",
            "word_extraction",
            "image_ocr",
            "table_extraction",
            "financial_statement_parsing",
            "ratio_extraction",
            "data_validation",
            "format_conversion",
            "arabic_text_processing"
        ]

        self.state.capabilities.extend(extraction_capabilities)

        # Set specializations based on agent name
        if "ocr" in self.state.agent_id.lower():
            self.state.specializations.extend(["ocr_processing", "image_extraction", "scanned_documents"])
        elif "validator" in self.state.agent_id.lower():
            self.state.specializations.extend(["data_validation", "quality_assurance", "accuracy_checking"])
        else:
            self.state.specializations.extend(["general_extraction", "multi_format_processing"])

    def _initialize_prompts(self) -> None:
        """Initialize data extraction prompts"""
        super()._initialize_prompts()

        self.extraction_prompt = ChatPromptTemplate.from_messages([
            ("system", """
            أنت خبير في استخراج البيانات المالية من المستندات المختلفة.
            مهمتك هي تحديد وتنظيم البيانات المالية الرئيسية من النصوص المقدمة.

            You are an expert in extracting financial data from various documents.
            Your task is to identify and organize key financial data from the provided text.

            قم بالتركيز على:
            - القوائم المالية (قائمة المركز المالي، قائمة الدخل، قائمة التدفقات النقدية)
            - النسب المالية والمؤشرات
            - البيانات التاريخية والمقارنات
            - معلومات الشركة الأساسية

            Focus on:
            - Financial statements (Balance Sheet, Income Statement, Cash Flow Statement)
            - Financial ratios and indicators
            - Historical data and comparisons
            - Basic company information
            """),
            ("human", """
            استخرج البيانات المالية من النص التالي:

            {document_text}

            Extract financial data from the following text:

            قم بتنظيم البيانات في تنسيق JSON منظم يتضمن:
            1. معلومات الشركة الأساسية
            2. القوائم المالية
            3. النسب المالية (إن وجدت)
            4. البيانات التاريخية (إن وجدت)

            Organize the data in a structured JSON format including:
            1. Basic company information
            2. Financial statements
            3. Financial ratios (if available)
            4. Historical data (if available)
            """)
        ])

        self.ocr_enhancement_prompt = ChatPromptTemplate.from_messages([
            ("system", """
            أنت خبير في تحسين وتصحيح النصوص المستخرجة من OCR.
            مهمتك هي تصحيح الأخطاء وتحسين جودة النص المالي المستخرج.

            You are an expert in enhancing and correcting OCR-extracted text.
            Your task is to correct errors and improve the quality of extracted financial text.

            قم بالتركيز على:
            - تصحيح الأرقام المالية
            - تصحيح أسماء الحسابات المالية
            - تحسين تنسيق الجداول
            - التعامل مع النصوص العربية والإنجليزية

            Focus on:
            - Correcting financial numbers
            - Correcting financial account names
            - Improving table formatting
            - Handling Arabic and English text
            """),
            ("human", """
            قم بتحسين وتصحيح النص التالي المستخرج من OCR:

            {ocr_text}

            Enhance and correct the following OCR-extracted text:

            قدم النص المحسن مع تصحيح الأخطاء وتحسين التنسيق.
            Provide the enhanced text with error corrections and improved formatting.
            """)
        ])

    def _initialize_extraction_patterns(self) -> Dict[str, Any]:
        """Initialize financial data extraction patterns"""
        return {
            "financial_statements": {
                "balance_sheet": {
                    "arabic_headers": [
                        "قائمة المركز المالي", "الميزانية العمومية", "قائمة الأصول والخصوم",
                        "المركز المالي", "قائمة المركز المالى"
                    ],
                    "english_headers": [
                        "balance sheet", "statement of financial position", "assets and liabilities",
                        "statement of position"
                    ],
                    "accounts": {
                        "assets": {
                            "arabic": ["الأصول", "إجمالي الأصول", "الأصول المتداولة", "الأصول غير المتداولة"],
                            "english": ["assets", "total assets", "current assets", "non-current assets"]
                        },
                        "liabilities": {
                            "arabic": ["الخصوم", "إجمالي الخصوم", "الخصوم المتداولة", "الخصوم غير المتداولة"],
                            "english": ["liabilities", "total liabilities", "current liabilities", "non-current liabilities"]
                        },
                        "equity": {
                            "arabic": ["حقوق الملكية", "حقوق المساهمين", "رأس المال"],
                            "english": ["equity", "shareholders' equity", "capital", "retained earnings"]
                        }
                    }
                },
                "income_statement": {
                    "arabic_headers": [
                        "قائمة الدخل", "قائمة الأرباح والخسائر", "قائمة الدخل الشامل",
                        "بيان الدخل", "قائمة الايرادات والمصروفات"
                    ],
                    "english_headers": [
                        "income statement", "profit and loss", "statement of comprehensive income",
                        "statement of operations", "earnings statement"
                    ],
                    "accounts": {
                        "revenue": {
                            "arabic": ["الإيرادات", "المبيعات", "إجمالي الإيرادات", "الدخل"],
                            "english": ["revenue", "sales", "total revenue", "income", "turnover"]
                        },
                        "expenses": {
                            "arabic": ["المصروفات", "التكاليف", "تكلفة المبيعات", "المصاريف"],
                            "english": ["expenses", "costs", "cost of sales", "operating expenses"]
                        },
                        "profit": {
                            "arabic": ["الربح", "صافي الربح", "الربح التشغيلي", "الأرباح"],
                            "english": ["profit", "net profit", "operating profit", "earnings", "net income"]
                        }
                    }
                },
                "cash_flow": {
                    "arabic_headers": [
                        "قائمة التدفقات النقدية", "بيان التدفق النقدي", "قائمة النقدية",
                        "التدفقات النقدية"
                    ],
                    "english_headers": [
                        "cash flow statement", "statement of cash flows", "cash flows"
                    ],
                    "sections": {
                        "operating": {
                            "arabic": ["الأنشطة التشغيلية", "التدفق النقدي التشغيلي"],
                            "english": ["operating activities", "operating cash flow"]
                        },
                        "investing": {
                            "arabic": ["الأنشطة الاستثمارية", "التدفق النقدي الاستثماري"],
                            "english": ["investing activities", "investing cash flow"]
                        },
                        "financing": {
                            "arabic": ["الأنشطة التمويلية", "التدفق النقدي التمويلي"],
                            "english": ["financing activities", "financing cash flow"]
                        }
                    }
                }
            },
            "financial_ratios": {
                "liquidity": {
                    "arabic": ["نسبة التداول", "نسبة السيولة", "نسبة السيولة السريعة"],
                    "english": ["current ratio", "liquidity ratio", "quick ratio", "acid test"]
                },
                "profitability": {
                    "arabic": ["هامش الربح", "العائد على الأصول", "العائد على حقوق الملكية"],
                    "english": ["profit margin", "return on assets", "return on equity", "ROA", "ROE"]
                },
                "leverage": {
                    "arabic": ["نسبة الدين", "نسبة الدين إلى حقوق الملكية", "الرافعة المالية"],
                    "english": ["debt ratio", "debt to equity", "leverage ratio", "financial leverage"]
                }
            },
            "currency_patterns": {
                "sar": ["ريال", "ر.س", "SAR", "SR"],
                "usd": ["دولار", "USD", "$"],
                "aed": ["درهم", "AED", "DH"],
                "eur": ["يورو", "EUR", "€"]
            },
            "number_patterns": {
                "arabic_numbers": r'[\u0660-\u0669]+',  # Arabic-Indic digits
                "english_numbers": r'\d+(?:,\d{3})*(?:\.\d+)?',
                "currency_amounts": r'[\d,]+\.?\d*\s*(?:ريال|دولار|درهم|SAR|USD|AED|SR|\$)',
                "percentages": r'\d+\.?\d*\s*%'
            }
        }

    def _initialize_document_processors(self) -> Dict[str, Any]:
        """Initialize document format processors"""
        return {
            "pdf": self._process_pdf,
            "excel": self._process_excel,
            "word": self._process_word,
            "image": self._process_image,
            "csv": self._process_csv
        }

    async def process_task(self, task: AgentTask) -> Dict[str, Any]:
        """Process data extraction task"""
        task_type = task.task_type
        input_data = task.input_data

        if task_type == "data_extraction":
            return await self._extract_from_documents(input_data)
        elif task_type == "quick_data_extraction":
            return await self._quick_extract_from_documents(input_data)
        elif task_type == "ocr_processing":
            return await self._process_ocr_documents(input_data)
        elif task_type == "data_validation":
            return await self._validate_extracted_data(input_data)
        elif task_type == "format_conversion":
            return await self._convert_data_format(input_data)
        else:
            raise ValueError(f"Unsupported task type: {task_type}")

    async def _extract_from_documents(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """Extract financial data from uploaded documents"""
        documents = input_data.get("documents", [])
        extraction_results = []

        for doc in documents:
            try:
                doc_result = await self._process_single_document(doc)
                extraction_results.append(doc_result)
            except Exception as e:
                self.logger.error(f"Failed to process document {doc.get('name', 'unknown')}: {str(e)}")
                extraction_results.append({
                    "document_name": doc.get("name", "unknown"),
                    "status": "failed",
                    "error": str(e)
                })

        # Consolidate results
        consolidated_data = await self._consolidate_extraction_results(extraction_results)

        return {
            "status": "completed",
            "extracted_data": consolidated_data,
            "individual_results": extraction_results,
            "processing_metadata": {
                "total_documents": len(documents),
                "successful_extractions": len([r for r in extraction_results if r.get("status") == "success"]),
                "extraction_confidence": self._calculate_extraction_confidence(extraction_results)
            }
        }

    async def _quick_extract_from_documents(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """Quick extraction for urgent analysis (simplified processing)"""
        documents = input_data.get("documents", [])

        # Process only the first document or most relevant document
        if not documents:
            return {"status": "failed", "error": "No documents provided"}

        # Select best document for quick processing
        selected_doc = self._select_best_document_for_quick_processing(documents)

        try:
            doc_result = await self._process_single_document(selected_doc, quick_mode=True)

            # Extract only essential financial metrics
            essential_data = self._extract_essential_metrics(doc_result.get("extracted_data", {}))

            return {
                "status": "completed",
                "extracted_data": essential_data,
                "source_document": selected_doc.get("name", "unknown"),
                "processing_metadata": {
                    "quick_mode": True,
                    "processing_time": "reduced",
                    "confidence": doc_result.get("confidence", 0.8)
                }
            }

        except Exception as e:
            return {
                "status": "failed",
                "error": str(e),
                "source_document": selected_doc.get("name", "unknown")
            }

    async def _process_single_document(self, document: Dict[str, Any], quick_mode: bool = False) -> Dict[str, Any]:
        """Process a single document for data extraction"""
        doc_name = document.get("name", "")
        doc_type = document.get("type", "")
        doc_content = document.get("content", "")

        # Determine processing method based on file type
        if doc_type.lower().endswith('.pdf') or 'pdf' in doc_type.lower():
            text_content = await self._process_pdf(doc_content)
        elif doc_type.lower().endswith(('.xlsx', '.xls')) or 'excel' in doc_type.lower():
            text_content = await self._process_excel(doc_content)
        elif doc_type.lower().endswith('.docx') or 'word' in doc_type.lower():
            text_content = await self._process_word(doc_content)
        elif doc_type.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff')) or 'image' in doc_type.lower():
            text_content = await self._process_image(doc_content)
        elif doc_type.lower().endswith('.csv') or 'csv' in doc_type.lower():
            text_content = await self._process_csv(doc_content)
        else:
            # Try to process as text
            text_content = doc_content

        # Extract financial data using AI
        if quick_mode:
            extracted_data = await self._quick_extract_financial_data(text_content)
        else:
            extracted_data = await self._extract_financial_data_with_ai(text_content)

        return {
            "document_name": doc_name,
            "document_type": doc_type,
            "status": "success",
            "extracted_data": extracted_data,
            "confidence": self._calculate_confidence(extracted_data, text_content),
            "raw_text_length": len(text_content)
        }

    async def _process_pdf(self, pdf_content: str) -> str:
        """Extract text from PDF content"""
        try:
            # If content is base64 encoded
            if isinstance(pdf_content, str) and pdf_content.startswith('data:'):
                # Remove data URL prefix
                pdf_content = pdf_content.split(',')[1]
                pdf_bytes = base64.b64decode(pdf_content)
            else:
                pdf_bytes = pdf_content.encode() if isinstance(pdf_content, str) else pdf_content

            # Use PyPDF2 to extract text
            pdf_file = io.BytesIO(pdf_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_content = ""
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"

            return text_content

        except Exception as e:
            self.logger.error(f"PDF processing failed: {str(e)}")
            return ""

    async def _process_excel(self, excel_content: str) -> str:
        """Extract data from Excel content"""
        try:
            # If content is base64 encoded
            if isinstance(excel_content, str) and excel_content.startswith('data:'):
                excel_content = excel_content.split(',')[1]
                excel_bytes = base64.b64decode(excel_content)
            else:
                excel_bytes = excel_content.encode() if isinstance(excel_content, str) else excel_content

            # Use pandas to read Excel
            excel_file = io.BytesIO(excel_bytes)
            workbook = pd.read_excel(excel_file, sheet_name=None)  # Read all sheets

            text_content = ""
            for sheet_name, df in workbook.items():
                text_content += f"Sheet: {sheet_name}\n"
                text_content += df.to_string() + "\n\n"

            return text_content

        except Exception as e:
            self.logger.error(f"Excel processing failed: {str(e)}")
            return ""

    async def _process_word(self, word_content: str) -> str:
        """Extract text from Word document content"""
        try:
            # If content is base64 encoded
            if isinstance(word_content, str) and word_content.startswith('data:'):
                word_content = word_content.split(',')[1]
                word_bytes = base64.b64decode(word_content)
            else:
                word_bytes = word_content.encode() if isinstance(word_content, str) else word_content

            # Use python-docx to extract text
            word_file = io.BytesIO(word_bytes)
            doc = Document(word_file)

            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + "\t"
                    text_content += "\n"

            return text_content

        except Exception as e:
            self.logger.error(f"Word processing failed: {str(e)}")
            return ""

    async def _process_image(self, image_content: str) -> str:
        """Extract text from image using OCR"""
        try:
            # If content is base64 encoded
            if isinstance(image_content, str) and image_content.startswith('data:'):
                image_content = image_content.split(',')[1]
                image_bytes = base64.b64decode(image_content)
            else:
                image_bytes = image_content.encode() if isinstance(image_content, str) else image_content

            # Use PIL and pytesseract for OCR
            image = Image.open(io.BytesIO(image_bytes))

            # OCR with Arabic and English support
            text_content = pytesseract.image_to_string(
                image,
                lang='ara+eng',  # Arabic + English
                config='--oem 3 --psm 6'
            )

            # Enhance OCR results using AI
            enhanced_text = await self._enhance_ocr_text(text_content)

            return enhanced_text

        except Exception as e:
            self.logger.error(f"Image OCR processing failed: {str(e)}")
            return ""

    async def _process_csv(self, csv_content: str) -> str:
        """Process CSV content"""
        try:
            # If content is base64 encoded
            if isinstance(csv_content, str) and csv_content.startswith('data:'):
                csv_content = csv_content.split(',')[1]
                csv_bytes = base64.b64decode(csv_content)
                csv_content = csv_bytes.decode('utf-8')

            # Use pandas to read CSV
            csv_file = io.StringIO(csv_content)
            df = pd.read_csv(csv_file)

            text_content = df.to_string()
            return text_content

        except Exception as e:
            self.logger.error(f"CSV processing failed: {str(e)}")
            return csv_content  # Return as-is if processing fails

    async def _enhance_ocr_text(self, ocr_text: str) -> str:
        """Enhance OCR text using AI to correct errors"""
        try:
            chain = self.ocr_enhancement_prompt | self.llm
            response = await chain.ainvoke({"ocr_text": ocr_text})
            return response.content
        except Exception as e:
            self.logger.error(f"OCR enhancement failed: {str(e)}")
            return ocr_text  # Return original if enhancement fails

    async def _extract_financial_data_with_ai(self, text_content: str) -> Dict[str, Any]:
        """Extract financial data using AI analysis"""
        try:
            chain = self.extraction_prompt | self.llm
            response = await chain.ainvoke({"document_text": text_content})

            # Try to parse JSON response
            try:
                extracted_data = json.loads(response.content)
            except json.JSONDecodeError:
                # If not JSON, structure the response
                extracted_data = {
                    "raw_analysis": response.content,
                    "structured_data": self._parse_unstructured_response(response.content)
                }

            # Enhance with pattern-based extraction
            pattern_extracted = await self._extract_using_patterns(text_content)
            extracted_data["pattern_extracted"] = pattern_extracted

            return extracted_data

        except Exception as e:
            self.logger.error(f"AI extraction failed: {str(e)}")
            # Fallback to pattern-based extraction only
            return await self._extract_using_patterns(text_content)

    async def _quick_extract_financial_data(self, text_content: str) -> Dict[str, Any]:
        """Quick financial data extraction for urgent analysis"""
        # Focus on key metrics only
        pattern_data = await self._extract_using_patterns(text_content)

        # Extract essential metrics
        essential_data = {
            "company_info": self._extract_company_info(text_content),
            "key_figures": self._extract_key_financial_figures(text_content),
            "basic_ratios": pattern_data.get("financial_ratios", {}),
            "extraction_method": "quick_pattern_based"
        }

        return essential_data

    async def _extract_using_patterns(self, text_content: str) -> Dict[str, Any]:
        """Extract financial data using predefined patterns"""
        extracted_data = {
            "balance_sheet": {},
            "income_statement": {},
            "cash_flow_statement": {},
            "financial_ratios": {},
            "company_info": {},
            "currency": "SAR"  # Default
        }

        # Extract company information
        extracted_data["company_info"] = self._extract_company_info(text_content)

        # Extract financial statements
        extracted_data["balance_sheet"] = self._extract_balance_sheet_data(text_content)
        extracted_data["income_statement"] = self._extract_income_statement_data(text_content)
        extracted_data["cash_flow_statement"] = self._extract_cash_flow_data(text_content)

        # Extract financial ratios
        extracted_data["financial_ratios"] = self._extract_financial_ratios(text_content)

        # Detect currency
        extracted_data["currency"] = self._detect_currency(text_content)

        return extracted_data

    def _extract_company_info(self, text: str) -> Dict[str, Any]:
        """Extract basic company information"""
        company_info = {}

        # Company name patterns
        name_patterns = [
            r'شركة\s+([^\n\r]+)',
            r'Company\s+([^\n\r]+)',
            r'شركة\s+([^،]+)',
            r'([^\n\r]+)\s+Company',
            r'([^\n\r]+)\s+Corporation',
            r'([^\n\r]+)\s+المحدودة'
        ]

        for pattern in name_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                company_info["name"] = match.group(1).strip()
                break

        # Extract sector/industry
        sector_keywords = {
            "banking": ["بنك", "مصرفية", "bank", "banking"],
            "energy": ["طاقة", "نفط", "energy", "oil", "petroleum"],
            "telecom": ["اتصالات", "telecommunications", "telecom"],
            "retail": ["تجزئة", "retail", "shopping"],
            "manufacturing": ["تصنيع", "manufacturing", "industrial"]
        }

        for sector, keywords in sector_keywords.items():
            for keyword in keywords:
                if keyword in text.lower():
                    company_info["sector"] = sector
                    break

        return company_info

    def _extract_balance_sheet_data(self, text: str) -> Dict[str, Any]:
        """Extract balance sheet data using patterns"""
        balance_sheet = {}

        # Assets patterns
        assets_patterns = self.financial_patterns["financial_statements"]["balance_sheet"]["accounts"]["assets"]

        for lang in ["arabic", "english"]:
            for term in assets_patterns[lang]:
                # Look for amounts following these terms
                pattern = rf'{re.escape(term)}\s*[:\-]?\s*([\d,]+\.?\d*)'
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    # Take the largest amount found (likely total assets)
                    amounts = [float(m.replace(',', '')) for m in matches]
                    balance_sheet["total_assets"] = max(amounts)
                    break

        # Similar logic for liabilities and equity
        liabilities_patterns = self.financial_patterns["financial_statements"]["balance_sheet"]["accounts"]["liabilities"]

        for lang in ["arabic", "english"]:
            for term in liabilities_patterns[lang]:
                pattern = rf'{re.escape(term)}\s*[:\-]?\s*([\d,]+\.?\d*)'
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    amounts = [float(m.replace(',', '')) for m in matches]
                    balance_sheet["total_liabilities"] = max(amounts)
                    break

        return balance_sheet

    def _extract_income_statement_data(self, text: str) -> Dict[str, Any]:
        """Extract income statement data using patterns"""
        income_statement = {}

        # Revenue patterns
        revenue_patterns = self.financial_patterns["financial_statements"]["income_statement"]["accounts"]["revenue"]

        for lang in ["arabic", "english"]:
            for term in revenue_patterns[lang]:
                pattern = rf'{re.escape(term)}\s*[:\-]?\s*([\d,]+\.?\d*)'
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    amounts = [float(m.replace(',', '')) for m in matches]
                    income_statement["revenue"] = max(amounts)
                    break

        # Profit patterns
        profit_patterns = self.financial_patterns["financial_statements"]["income_statement"]["accounts"]["profit"]

        for lang in ["arabic", "english"]:
            for term in profit_patterns[lang]:
                pattern = rf'{re.escape(term)}\s*[:\-]?\s*([\d,]+\.?\d*)'
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    amounts = [float(m.replace(',', '')) for m in matches]
                    income_statement["net_income"] = max(amounts)
                    break

        return income_statement

    def _extract_cash_flow_data(self, text: str) -> Dict[str, Any]:
        """Extract cash flow statement data using patterns"""
        cash_flow = {}

        # Operating cash flow patterns
        operating_patterns = self.financial_patterns["financial_statements"]["cash_flow"]["sections"]["operating"]

        for lang in ["arabic", "english"]:
            for term in operating_patterns[lang]:
                pattern = rf'{re.escape(term)}\s*[:\-]?\s*([\d,]+\.?\d*)'
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    amounts = [float(m.replace(',', '')) for m in matches]
                    cash_flow["operating_cash_flow"] = max(amounts)
                    break

        return cash_flow

    def _extract_financial_ratios(self, text: str) -> Dict[str, Any]:
        """Extract financial ratios using patterns"""
        ratios = {}

        # Liquidity ratios
        liquidity_patterns = self.financial_patterns["financial_ratios"]["liquidity"]

        for lang in ["arabic", "english"]:
            for term in liquidity_patterns[lang]:
                pattern = rf'{re.escape(term)}\s*[:\-]?\s*([\d.]+)'
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    ratios["current_ratio"] = float(matches[0])
                    break

        return ratios

    def _extract_key_financial_figures(self, text: str) -> Dict[str, Any]:
        """Extract key financial figures for quick analysis"""
        key_figures = {}

        # Look for major financial numbers
        number_pattern = r'([\d,]+\.?\d*)\s*(?:ريال|ألف|مليون|billion|million|thousand)'
        numbers = re.findall(number_pattern, text, re.IGNORECASE)

        if numbers:
            # Convert and sort numbers
            amounts = []
            for num_str in numbers:
                try:
                    amount = float(num_str.replace(',', ''))
                    amounts.append(amount)
                except ValueError:
                    continue

            amounts.sort(reverse=True)

            # Assign largest amounts to likely categories
            if len(amounts) >= 1:
                key_figures["largest_amount"] = amounts[0]
            if len(amounts) >= 2:
                key_figures["second_largest_amount"] = amounts[1]
            if len(amounts) >= 3:
                key_figures["third_largest_amount"] = amounts[2]

        return key_figures

    def _detect_currency(self, text: str) -> str:
        """Detect the currency used in the document"""
        currency_patterns = self.financial_patterns["currency_patterns"]

        for currency, patterns in currency_patterns.items():
            for pattern in patterns:
                if pattern in text:
                    return currency.upper()

        return "SAR"  # Default to Saudi Riyal

    def _parse_unstructured_response(self, response: str) -> Dict[str, Any]:
        """Parse unstructured AI response into structured data"""
        # Simple parsing logic - would be more sophisticated in production
        structured = {
            "analysis_available": True,
            "key_points": []
        }

        # Extract bullet points or numbered lists
        points = re.findall(r'[•\-\*]\s*([^\n\r]+)', response)
        structured["key_points"] = points

        return structured

    def _calculate_confidence(self, extracted_data: Dict[str, Any], original_text: str) -> float:
        """Calculate confidence score for extracted data"""
        confidence = 0.5  # Base confidence

        # Increase confidence based on data completeness
        if extracted_data.get("balance_sheet"):
            confidence += 0.1
        if extracted_data.get("income_statement"):
            confidence += 0.1
        if extracted_data.get("cash_flow_statement"):
            confidence += 0.1

        # Increase confidence based on company info
        if extracted_data.get("company_info", {}).get("name"):
            confidence += 0.1

        # Increase confidence based on financial ratios
        if extracted_data.get("financial_ratios"):
            confidence += 0.1

        # Decrease confidence for very short texts
        if len(original_text) < 500:
            confidence *= 0.8

        return min(1.0, max(0.1, confidence))

    def _select_best_document_for_quick_processing(self, documents: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Select the best document for quick processing"""
        # Prioritize by file type and size
        priority_types = ['excel', 'csv', 'pdf', 'word', 'image']

        for doc_type in priority_types:
            for doc in documents:
                if doc_type in doc.get('type', '').lower():
                    return doc

        # Return first document if no preferred type found
        return documents[0] if documents else {}

    def _extract_essential_metrics(self, extracted_data: Dict[str, Any]) -> Dict[str, Any]:
        """Extract only essential metrics for quick analysis"""
        essential = {
            "company_name": extracted_data.get("company_info", {}).get("name", "Unknown"),
            "total_assets": extracted_data.get("balance_sheet", {}).get("total_assets"),
            "total_revenue": extracted_data.get("income_statement", {}).get("revenue"),
            "net_income": extracted_data.get("income_statement", {}).get("net_income"),
            "operating_cash_flow": extracted_data.get("cash_flow_statement", {}).get("operating_cash_flow"),
            "current_ratio": extracted_data.get("financial_ratios", {}).get("current_ratio"),
            "currency": extracted_data.get("currency", "SAR")
        }

        # Remove None values
        essential = {k: v for k, v in essential.items() if v is not None}

        return essential

    async def _consolidate_extraction_results(self, extraction_results: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Consolidate results from multiple documents"""
        consolidated = {
            "company_info": {},
            "balance_sheet": {},
            "income_statement": {},
            "cash_flow_statement": {},
            "financial_ratios": {},
            "source_documents": []
        }

        successful_results = [r for r in extraction_results if r.get("status") == "success"]

        for result in successful_results:
            extracted = result.get("extracted_data", {})
            doc_name = result.get("document_name", "unknown")

            consolidated["source_documents"].append(doc_name)

            # Merge company info (prefer first complete info)
            if extracted.get("company_info") and not consolidated["company_info"].get("name"):
                consolidated["company_info"].update(extracted["company_info"])

            # Merge financial statements (prefer most complete data)
            for statement_type in ["balance_sheet", "income_statement", "cash_flow_statement"]:
                if extracted.get(statement_type):
                    if not consolidated[statement_type] or len(extracted[statement_type]) > len(consolidated[statement_type]):
                        consolidated[statement_type].update(extracted[statement_type])

            # Merge financial ratios
            if extracted.get("financial_ratios"):
                consolidated["financial_ratios"].update(extracted["financial_ratios"])

        return consolidated

    def _calculate_extraction_confidence(self, extraction_results: List[Dict[str, Any]]) -> float:
        """Calculate overall extraction confidence"""
        if not extraction_results:
            return 0.0

        successful_results = [r for r in extraction_results if r.get("status") == "success"]

        if not successful_results:
            return 0.0

        # Average confidence of successful extractions
        confidences = [r.get("confidence", 0.5) for r in successful_results]
        avg_confidence = sum(confidences) / len(confidences)

        # Adjust based on success rate
        success_rate = len(successful_results) / len(extraction_results)
        adjusted_confidence = avg_confidence * success_rate

        return adjusted_confidence

    async def _validate_extracted_data(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate extracted financial data for consistency and accuracy"""
        extracted_data = input_data.get("extracted_data", {})
        validation_results = {
            "is_valid": True,
            "validation_errors": [],
            "validation_warnings": [],
            "corrected_data": extracted_data.copy()
        }

        # Validate balance sheet equation
        balance_sheet = extracted_data.get("balance_sheet", {})
        if balance_sheet.get("total_assets") and balance_sheet.get("total_liabilities") and balance_sheet.get("total_equity"):
            assets = balance_sheet["total_assets"]
            liabilities = balance_sheet["total_liabilities"]
            equity = balance_sheet["total_equity"]

            if abs(assets - (liabilities + equity)) > assets * 0.05:  # 5% tolerance
                validation_results["validation_errors"].append(
                    "Balance sheet equation does not balance: Assets ≠ Liabilities + Equity"
                )
                validation_results["is_valid"] = False

        # Validate income statement consistency
        income_statement = extracted_data.get("income_statement", {})
        if income_statement.get("revenue") and income_statement.get("net_income"):
            if income_statement["net_income"] > income_statement["revenue"]:
                validation_results["validation_warnings"].append(
                    "Net income exceeds revenue - please verify"
                )

        # Validate financial ratios ranges
        ratios = extracted_data.get("financial_ratios", {})
        if ratios.get("current_ratio"):
            if ratios["current_ratio"] < 0 or ratios["current_ratio"] > 100:
                validation_results["validation_errors"].append(
                    f"Current ratio {ratios['current_ratio']} seems unrealistic"
                )
                validation_results["is_valid"] = False

        return validation_results

    async def _convert_data_format(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """Convert extracted data to different formats"""
        extracted_data = input_data.get("extracted_data", {})
        target_format = input_data.get("target_format", "standard")

        if target_format == "analysis_ready":
            # Convert to format suitable for financial analysis engine
            converted_data = {
                "financial_statements": {
                    "balance_sheet": extracted_data.get("balance_sheet", {}),
                    "income_statement": extracted_data.get("income_statement", {}),
                    "cash_flow_statement": extracted_data.get("cash_flow_statement", {})
                },
                "company_info": extracted_data.get("company_info", {}),
                "metadata": {
                    "currency": extracted_data.get("currency", "SAR"),
                    "source_documents": extracted_data.get("source_documents", []),
                    "extraction_date": datetime.now().isoformat()
                }
            }

        elif target_format == "json_export":
            # Convert to clean JSON format
            converted_data = {
                "company": extracted_data.get("company_info", {}),
                "financials": {
                    "balance_sheet": extracted_data.get("balance_sheet", {}),
                    "income_statement": extracted_data.get("income_statement", {}),
                    "cash_flow": extracted_data.get("cash_flow_statement", {})
                },
                "ratios": extracted_data.get("financial_ratios", {}),
                "metadata": {
                    "currency": extracted_data.get("currency", "SAR"),
                    "extraction_timestamp": datetime.now().isoformat()
                }
            }

        else:
            # Standard format
            converted_data = extracted_data

        return {
            "status": "completed",
            "converted_data": converted_data,
            "target_format": target_format
        }

    async def _process_ocr_documents(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """Specialized OCR processing for scanned documents"""
        documents = input_data.get("documents", [])
        ocr_results = []

        for doc in documents:
            if doc.get("type", "").lower() in ['image/png', 'image/jpeg', 'image/jpg', 'image/tiff']:
                try:
                    # Enhanced OCR processing
                    text_content = await self._process_image(doc.get("content", ""))

                    # Post-process OCR text
                    cleaned_text = await self._post_process_ocr_text(text_content)

                    # Extract financial data
                    extracted_data = await self._extract_using_patterns(cleaned_text)

                    ocr_results.append({
                        "document_name": doc.get("name", "unknown"),
                        "status": "success",
                        "raw_ocr_text": text_content,
                        "cleaned_text": cleaned_text,
                        "extracted_data": extracted_data,
                        "confidence": self._calculate_ocr_confidence(text_content, extracted_data)
                    })

                except Exception as e:
                    ocr_results.append({
                        "document_name": doc.get("name", "unknown"),
                        "status": "failed",
                        "error": str(e)
                    })

        return {
            "status": "completed",
            "ocr_results": ocr_results,
            "processing_metadata": {
                "total_images": len([d for d in documents if "image" in d.get("type", "").lower()]),
                "successful_ocr": len([r for r in ocr_results if r.get("status") == "success"])
            }
        }

    async def _post_process_ocr_text(self, ocr_text: str) -> str:
        """Post-process OCR text to improve quality"""
        # Basic cleaning
        cleaned_text = ocr_text.strip()

        # Fix common OCR errors for financial documents
        replacements = {
            'O': '0',  # Letter O to zero
            'l': '1',  # Letter l to one
            'S': '5',  # In certain contexts
            '|': '1',  # Pipe to one
        }

        # Apply replacements in number contexts
        for old, new in replacements.items():
            # Only replace in number contexts
            cleaned_text = re.sub(rf'\b{old}(?=\d)', new, cleaned_text)
            cleaned_text = re.sub(rf'(?<=\d){old}\b', new, cleaned_text)

        return cleaned_text

    def _calculate_ocr_confidence(self, ocr_text: str, extracted_data: Dict[str, Any]) -> float:
        """Calculate confidence score for OCR results"""
        confidence = 0.6  # Base OCR confidence

        # Increase confidence based on text length and content
        if len(ocr_text) > 100:
            confidence += 0.1

        # Check for financial keywords
        financial_keywords = ['ريال', 'دولار', 'أصول', 'خصوم', 'ربح', 'إيرادات', 'assets', 'liabilities', 'revenue', 'profit']
        keyword_count = sum(1 for keyword in financial_keywords if keyword in ocr_text.lower())
        confidence += min(0.2, keyword_count * 0.05)

        # Increase confidence based on extracted data quality
        if extracted_data.get("balance_sheet"):
            confidence += 0.1
        if extracted_data.get("income_statement"):
            confidence += 0.1

        return min(1.0, max(0.1, confidence))